"""
evaluate.py — Quantitative validation for VisionAST
=====================================================
Produces the numbers a paper's Results section actually needs:

  1. Disc-level measurement accuracy vs. Dryad ground-truth mm
     (Essential Agreement, MAE, RMSE, correlation)
  2. S/I/R categorical agreement using CLSI/FDA-style error taxonomy
     (Categorical Agreement, Major Error, Very Major Error, minor error)
  3. Antibiotic-disc detection recall (Hough detector vs. ground-truth disc count)
  4. Image-quality-gate behaviour (how many held-out images get flagged/rejected)
  5. (optional) Pixel-level Dice/IoU against the auto-generated training masks,
     for a sanity check consistent with train.py's own metric

IMPORTANT CAVEAT (read before trusting the antibiotic-level numbers):
  The Dryad docx tables give (antibiotic, zone_mm) per sample but NOT the
  pixel location of each disc. Training-mask generation (generate_masks_dryad.py)
  assigned docx row i -> the i-th disc when discs are sorted LEFT-TO-RIGHT by
  x-coordinate. This script uses that same convention to match predicted discs
  back to ground-truth rows. It is the best available proxy, not verified
  ground-truth positional alignment. Treat per-antibiotic error type counts
  (ME/VME) as indicative, not certified, until a human spot-checks a sample
  of the matches (the script saves annotated images for exactly this reason).

Usage:
    python src/evaluate.py \
        --model_path  models/best_model.pth \
        --model_size  small \
        --images_dir  data/dryad/images_original \
        --tables_dir  data/dryad/Tables \
        --masks_dir   data/dryad/masks \
        --output_dir  eval_results/ \
        --test_frac   0.25 \
        --seed        42 \
        --tolerance_mm 3.0
"""

import argparse
import json
import re
import sys
import random
import csv
from pathlib import Path

import numpy as np
import cv2
import torch

sys.path.insert(0, str(Path(__file__).parent))
from model import build_model
from predict import (
    validate_image, crop_to_plate, find_discs, segment,
    measure_and_annotate, classify, BREAKPOINTS_EUCAST, DISC_DIAMETER_MM,
)

try:
    from docx import Document
except ImportError:
    raise ImportError("Run:  pip install python-docx")


# ══════════════════════════════════════════════════════════════════════════
# Dataset matching (same convention as generate_masks_dryad.py)
# ══════════════════════════════════════════════════════════════════════════

SUPPORTED_EXT = {".jpg", ".jpeg", ".png", ".tif", ".tiff"}


def parse_docx(docx_path: Path):
    """Returns list of (antibiotic_abbrev, zone_mm) from one Table docx."""
    doc = Document(str(docx_path))
    rows = []
    if not doc.tables:
        return rows
    tbl = doc.tables[0]
    for row in tbl.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 3:
            continue
        abbrev = cells[0]
        try:
            zone_mm = float(cells[2])
            if zone_mm > 0:
                rows.append((abbrev, zone_mm))
        except ValueError:
            continue
    return rows


def build_id_maps(images_dir: Path, tables_dir: Path):
    table_map = {}
    for docx in tables_dir.glob("*.docx"):
        m = re.search(r'Table\s+([\d\.]+)', docx.name)
        if m:
            table_map[m.group(1).rstrip('.').strip()] = docx

    image_map = {}
    for img in images_dir.iterdir():
        if img.suffix.lower() not in SUPPORTED_EXT:
            continue
        m = re.match(r'^([\d\.]+)\s*', img.name)
        if m:
            image_map[m.group(1).rstrip('.').strip()] = img

    common_ids = sorted(set(table_map) & set(image_map))
    return common_ids, image_map, table_map


def make_eval_split(common_ids, test_frac, seed, test_ids_file=None):
    if test_ids_file and Path(test_ids_file).exists():
        with open(test_ids_file) as f:
            ids = [line.strip() for line in f if line.strip()]
        missing = [i for i in ids if i not in common_ids]
        if missing:
            print(f"⚠️  {len(missing)} ids from {test_ids_file} not found in dataset — skipping them.")
        return [i for i in ids if i in common_ids]

    rng = random.Random(seed)
    ids = common_ids[:]
    rng.shuffle(ids)
    n_test = max(1, int(len(ids) * test_frac))
    return sorted(ids[:n_test])


# ══════════════════════════════════════════════════════════════════════════
# CLSI/FDA-style S/I/R error taxonomy
# ══════════════════════════════════════════════════════════════════════════

def error_type(gt_class, pred_class):
    if gt_class == pred_class:
        return "agree"
    if gt_class == "S" and pred_class == "R":
        return "ME"       # Major Error: false-resistant
    if gt_class == "R" and pred_class == "S":
        return "VME"      # Very Major Error: false-susceptible (most dangerous)
    return "minor"        # any mismatch involving "I"


# ══════════════════════════════════════════════════════════════════════════
# Pixel-level Dice/IoU (mirrors train.py's compute_metrics)
# ══════════════════════════════════════════════════════════════════════════

def pixel_metrics(pred_mask_bin: np.ndarray, gt_mask_bin: np.ndarray):
    p = (pred_mask_bin > 127).astype(np.uint8).flatten()
    g = (gt_mask_bin > 127).astype(np.uint8).flatten()
    tp = int(((p == 1) & (g == 1)).sum())
    fp = int(((p == 1) & (g == 0)).sum())
    fn = int(((p == 0) & (g == 1)).sum())
    tn = int(((p == 0) & (g == 0)).sum())
    eps = 1e-8
    return {
        "pixel_acc": (tp + tn) / (tp + fp + fn + tn + eps),
        "iou":       tp / (tp + fp + fn + eps),
        "dice":      2 * tp / (2 * tp + fp + fn + eps),
        "precision": tp / (tp + fp + eps),
        "recall":    tp / (tp + fn + eps),
    }


# ══════════════════════════════════════════════════════════════════════════
# Main evaluation loop
# ══════════════════════════════════════════════════════════════════════════

def main(args):
    device = "cuda" if torch.cuda.is_available() else "cpu"
    model = build_model(size=args.model_size, device=device, norm_type=args.norm)
    state = torch.load(args.model_path, map_location=device)
    if isinstance(state, dict) and "model_state" in state:
        state = state["model_state"]
    model.load_state_dict(state)
    model.eval()
    print(f"Model loaded | device={device}")

    images_dir = Path(args.images_dir)
    tables_dir = Path(args.tables_dir)
    masks_dir = Path(args.masks_dir) if args.masks_dir else None
    out_dir = Path(args.output_dir)
    (out_dir / "annotated").mkdir(parents=True, exist_ok=True)

    common_ids, image_map, table_map = build_id_maps(images_dir, tables_dir)
    print(f"Found {len(common_ids)} matched image/table pairs in {images_dir}")

    test_ids = make_eval_split(common_ids, args.test_frac, args.seed, args.test_ids_file)
    print(f"Evaluating on {len(test_ids)} held-out sample(s) "
          f"(seed={args.seed}, test_frac={args.test_frac})")

    # persist the split used, so results are reproducible / auditable
    split_file = out_dir / "test_split_ids.txt"
    with open(split_file, "w") as f:
        f.write("\n".join(test_ids))
    print(f"Test-set sample IDs saved to: {split_file}")

    disc_rows = []       # per-disc measurement comparisons
    image_rows = []      # per-image detection-recall / quality summary

    for sid in test_ids:
        img_path = image_map[sid]
        gt_zone_data = parse_docx(table_map[sid])
        if not gt_zone_data:
            continue

        img_bgr = cv2.imread(str(img_path))
        if img_bgr is None:
            print(f"⚠️  Could not read {img_path}, skipping.")
            continue

        # 1. quality gate — record behaviour but keep evaluating measurement
        #    accuracy regardless, so we can compare quality-passed vs flagged.
        qr = validate_image(img_bgr)
        errors = [i for i in qr.issues if i["severity"] == "error"]
        warns = [i for i in qr.issues if i["severity"] == "warning"]
        quality_status = "REJECTED" if errors else ("WARNED" if warns else "OK")

        # 2. crop, segment, detect discs (same pipeline as predict.py)
        if args.skip_crop:
            # DIAGNOSTIC: training (dataset.py) resizes the FULL raw image
            # directly with no crop_to_plate() step. Inference normally
            # crops to the plate first. That mismatch is a likely driver of
            # the Dice gap between training-time validation (~0.62) and
            # real inference (~0.32 measured here) — this flag bypasses the
            # crop so you can test that hypothesis on the existing
            # checkpoint before deciding whether to retrain anything.
            cropped_bgr, (ox, oy) = img_bgr, (0, 0)
        else:
            cropped_bgr, (ox, oy) = crop_to_plate(img_bgr)
        img_rgb = cv2.cvtColor(cropped_bgr, cv2.COLOR_BGR2RGB)
        gray = cv2.cvtColor(cropped_bgr, cv2.COLOR_BGR2GRAY)

        discs = find_discs(gray)
        # sort left-to-right to match the convention used when ground-truth
        # masks were generated from the docx tables (see module docstring)
        discs = sorted(discs, key=lambda c: c[0])
        # discs detected BEFORE segment() so the disc-scale-aware fragment
        # filter inside segment() can use them (see predict.py fix)
        mask = segment(model, img_rgb, args.img_size, device, discs=discs)

        n_gt = len(gt_zone_data)
        n_pred = len(discs)
        image_rows.append({
            "sample_id": sid, "image": img_path.name,
            "gt_disc_count": n_gt, "pred_disc_count": n_pred,
            "detection_recall": min(n_pred, n_gt) / n_gt if n_gt else None,
            "quality_status": quality_status,
            "quality_codes": ";".join(i["code"] for i in qr.issues),
        })

        if not discs:
            continue

        results, vis = measure_and_annotate(mask, discs, img_rgb, args.guideline, gray)
        cv2.imwrite(str(out_dir / "annotated" / f"{sid}_eval.jpg"), vis)

        # 3. pair predicted discs to ground-truth rows by index (both sorted
        #    left-to-right / docx row order) — see caveat in module docstring
        for i in range(min(len(results), n_gt)):
            gt_abbrev, gt_mm = gt_zone_data[i]
            pred_mm = results[i]["zone_diameter_mm"]
            no_zone = bool(results[i].get("note"))

            gt_class = classify(gt_mm, gt_abbrev, args.guideline)
            # classify the prediction using the GROUND-TRUTH antibiotic label,
            # not the (currently unreliable) OCR output — this isolates
            # measurement error from the separate disc-ID problem
            pred_class = classify(pred_mm, gt_abbrev, args.guideline)

            abs_err = abs(pred_mm - gt_mm)
            disc_rows.append({
                "sample_id": sid, "disc_index": i + 1,
                "antibiotic": gt_abbrev,
                "gt_mm": gt_mm, "pred_mm": pred_mm, "abs_error_mm": round(abs_err, 2),
                "within_tolerance": abs_err <= args.tolerance_mm,
                "gt_class": gt_class, "pred_class": pred_class,
                "error_type": error_type(gt_class, pred_class),
                "no_zone_detected": no_zone,
                "quality_status": quality_status,
            })

        # 4. optional pixel-level sanity check against training masks
        #    NOTE: `mask` is sized to the CROPPED plate region (from
        #    crop_to_plate), while the ground-truth masks in `masks_dir`
        #    were generated on the FULL original image. Comparing shapes
        #    directly silently skipped ~all images before this fix. Fix:
        #    paste the predicted mask back into a full-size canvas at its
        #    original crop offset (ox, oy) so both are in the same frame.
        if masks_dir:
            gt_mask_path = masks_dir / f"{img_path.stem}_mask.png"
            if gt_mask_path.exists():
                gt_mask = cv2.imread(str(gt_mask_path), cv2.IMREAD_GRAYSCALE)
                if gt_mask is not None:
                    full_h, full_w = img_bgr.shape[:2]
                    uncropped = np.zeros((full_h, full_w), dtype=np.uint8)
                    ch, cw = mask.shape[:2]
                    y2, x2 = min(oy + ch, full_h), min(ox + cw, full_w)
                    uncropped[oy:y2, ox:x2] = mask[:y2-oy, :x2-ox]
                    if uncropped.shape == gt_mask.shape:
                        pm = pixel_metrics(uncropped, gt_mask)
                        image_rows[-1].update({f"px_{k}": round(v, 4) for k, v in pm.items()})
                    else:
                        print(f"⚠️  {sid}: gt_mask shape {gt_mask.shape} still "
                              f"doesn't match image shape {(full_h, full_w)} — skipping pixel metrics.")

    if not disc_rows:
        print("❌  No disc-level comparisons produced — check paths/ids.")
        return

    # ── write per-disc + per-image CSVs ────────────────────────────────────
    disc_csv = out_dir / "per_disc_results.csv"
    with open(disc_csv, "w", newline="") as f:
        w = csv.DictWriter(f, fieldnames=list(disc_rows[0].keys()))
        w.writeheader(); w.writerows(disc_rows)

    image_csv = out_dir / "per_image_summary.csv"
    all_img_keys = sorted({k for r in image_rows for k in r.keys()})
    with open(image_csv, "w", newline="") as f:
        w = csv.DictWriter(f, fieldnames=all_img_keys)
        w.writeheader(); w.writerows(image_rows)

    # ── aggregate metrics ───────────────────────────────────────────────────
    gt_mm = np.array([r["gt_mm"] for r in disc_rows])
    pred_mm = np.array([r["pred_mm"] for r in disc_rows])
    abs_err = np.abs(pred_mm - gt_mm)

    n = len(disc_rows)
    ea_rate = float(np.mean(abs_err <= args.tolerance_mm))
    ca_rate = float(np.mean([r["gt_class"] == r["pred_class"] for r in disc_rows]))
    mae = float(np.mean(abs_err))
    rmse = float(np.sqrt(np.mean(abs_err ** 2)))
    corr = float(np.corrcoef(gt_mm, pred_mm)[0, 1]) if n > 1 else float("nan")
    no_zone_rate = float(np.mean([r["no_zone_detected"] for r in disc_rows]))

    n_true_S = sum(1 for r in disc_rows if r["gt_class"] == "S")
    n_true_R = sum(1 for r in disc_rows if r["gt_class"] == "R")
    n_ME = sum(1 for r in disc_rows if r["error_type"] == "ME")
    n_VME = sum(1 for r in disc_rows if r["error_type"] == "VME")
    me_rate = n_ME / n_true_S if n_true_S else float("nan")
    vme_rate = n_VME / n_true_R if n_true_R else float("nan")

    recalls = [r["detection_recall"] for r in image_rows if r["detection_recall"] is not None]
    mean_recall = float(np.mean(recalls)) if recalls else float("nan")

    q_counts = {}
    for r in image_rows:
        q_counts[r["quality_status"]] = q_counts.get(r["quality_status"], 0) + 1

    summary = {
        "n_images_evaluated": len(image_rows),
        "n_disc_comparisons": n,
        "tolerance_mm": args.tolerance_mm,
        "essential_agreement_pct": round(ea_rate * 100, 2),
        "categorical_agreement_pct": round(ca_rate * 100, 2),
        "mae_mm": round(mae, 3),
        "rmse_mm": round(rmse, 3),
        "pearson_r": round(corr, 3) if not np.isnan(corr) else None,
        "no_zone_detected_rate_pct": round(no_zone_rate * 100, 2),
        "major_error_rate_pct": round(me_rate * 100, 2) if not np.isnan(me_rate) else None,
        "very_major_error_rate_pct": round(vme_rate * 100, 2) if not np.isnan(vme_rate) else None,
        "n_true_S": n_true_S, "n_true_R": n_true_R,
        "n_ME": n_ME, "n_VME": n_VME,
        "mean_disc_detection_recall_pct": round(mean_recall * 100, 2) if not np.isnan(mean_recall) else None,
        "quality_gate_status_counts": q_counts,
    }

    summary_path = out_dir / "summary_metrics.json"
    with open(summary_path, "w") as f:
        json.dump(summary, f, indent=2)

    print("\n══════════════ EVALUATION SUMMARY ══════════════")
    for k, v in summary.items():
        print(f"  {k:>32}: {v}")
    print("══════════════════════════════════════════════════")
    print(f"\nPer-disc CSV   : {disc_csv}")
    print(f"Per-image CSV  : {image_csv}")
    print(f"Summary JSON   : {summary_path}")
    print(f"Annotated imgs : {out_dir/'annotated'}/  (spot-check disc/antibiotic matching here)")


if __name__ == "__main__":
    p = argparse.ArgumentParser(description="Quantitative validation for VisionAST")
    p.add_argument("--model_path", required=True)
    p.add_argument("--model_size", choices=["full", "small"], default="small")
    p.add_argument("--norm", choices=["batch","group"], default="batch",
                    help="MUST match the value used when the checkpoint was trained")
    p.add_argument("--img_size", type=int, default=512)
    p.add_argument("--guideline", choices=["EUCAST"], default="EUCAST")
    p.add_argument("--images_dir", default="data/dryad/images_original")
    p.add_argument("--tables_dir", default="data/dryad/Tables")
    p.add_argument("--masks_dir", default="data/dryad/masks",
                    help="Optional: for pixel-level Dice/IoU sanity check. Pass '' to skip.")
    p.add_argument("--output_dir", default="eval_results/")
    p.add_argument("--test_frac", type=float, default=0.25)
    p.add_argument("--seed", type=int, default=42)
    p.add_argument("--tolerance_mm", type=float, default=3.0,
                    help="Essential Agreement tolerance in mm (CLSI/FDA convention: 3mm)")
    p.add_argument("--test_ids_file", default=None,
                    help="Optional file with one sample id per line, to reuse a fixed test set")
    p.add_argument("--skip_crop", action="store_true",
                    help="Diagnostic: bypass crop_to_plate() to test the train/inference "
                         "preprocessing-mismatch hypothesis (see comment in main loop)")
    args = p.parse_args()
    if args.masks_dir == "":
        args.masks_dir = None
    main(args)