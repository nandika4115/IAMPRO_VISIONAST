"""
generate_masks_dryad.py  (UPDATED for actual Dryad dataset structure)
======================================================================
The Dryad dataset contains:
  - images_original/  :  "1.1.1. original.jpg", "1.2.1. original.jpg", ...
  - Tables/           :  "Table 1.1.1..docx", "Table 1.2.1..docx", ...

Each docx has one table with columns:
  Abbreviation | Antibiotic | Inhibition zone (mm) | Interpretation | Resistant | Susceptible

This script:
  1. Reads every docx to get zone diameters for each sample ID
  2. Finds the corresponding image
  3. Uses Hough circles to locate antibiotic discs on the plate
  4. Converts zone diameter (mm) → radius (px) using 6mm disc as ruler
  5. Draws and saves a binary mask

Usage:
    python src/generate_masks_dryad.py \
        --images_dir  data/dryad/images_original \
        --tables_dir  data/dryad/Tables \
        --masks_dir   data/dryad/masks \
        --visualise
"""

import os, re, math, argparse
import numpy as np
import cv2
from pathlib import Path
from tqdm import tqdm

try:
    from docx import Document
except ImportError:
    raise ImportError("Run:  pip install python-docx")

DISC_DIAMETER_MM  = 6.0
DISC_MIN_PX       = 8
DISC_MAX_PX       = 45
SUPPORTED_EXT     = {".jpg", ".jpeg", ".png", ".tif", ".tiff"}


# ── parse one docx → list of zone diameters (mm) ──────────────────────────────
def parse_docx(docx_path: Path) -> list:
    """Returns list of (antibiotic_abbrev, zone_mm) from one Table docx."""
    doc    = Document(str(docx_path))
    rows   = []
    if not doc.tables:
        return rows
    tbl = doc.tables[0]
    for row in tbl.rows[1:]:          # skip header row
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 3:
            continue
        abbrev   = cells[0]
        zone_str = cells[2]
        try:
            zone_mm = float(zone_str)
            if zone_mm > 0:
                rows.append((abbrev, zone_mm))
        except ValueError:
            continue
    return rows


# ── detect antibiotic discs via Hough ─────────────────────────────────────────
def find_discs(gray: np.ndarray):
    blurred = cv2.GaussianBlur(gray, (7, 7), 2)
    circles = cv2.HoughCircles(
        blurred, cv2.HOUGH_GRADIENT,
        dp=1, minDist=30, param1=60, param2=28,
        minRadius=DISC_MIN_PX, maxRadius=DISC_MAX_PX,
    )
    if circles is None:
        return []
    return [tuple(np.round(c).astype(int)) for c in circles[0]]


# ── draw one zone as a filled circle (halo only) ──────────────────────────────
def draw_zone(mask, cx, cy, zone_r_px, disc_r_px):
    cv2.circle(mask, (int(cx), int(cy)), int(zone_r_px), 255, -1)
    cv2.circle(mask, (int(cx), int(cy)), int(disc_r_px),   0, -1)


# ── main ──────────────────────────────────────────────────────────────────────
def main(images_dir, tables_dir, masks_dir, visualise):
    images_path = Path(images_dir)
    tables_path = Path(tables_dir)
    masks_path  = Path(masks_dir)
    masks_path.mkdir(parents=True, exist_ok=True)

    # build map:  sample_id (e.g. "1.1.1.") → docx path
    table_map = {}
    for docx in tables_path.glob("*.docx"):
        # "Table 1.1.1..docx"  →  "1.1.1"
        m = re.search(r'Table\s+([\d\.]+)', docx.name)
        if m:
            tbl_id = m.group(1).rstrip('.').strip()
            table_map[tbl_id] = docx

    # build map:  sample_id → image path
    image_map = {}
    for img in images_path.iterdir():
        if img.suffix.lower() not in SUPPORTED_EXT:
            continue
        # "1.1.1. original.jpg" → "1.1.1"
        m = re.match(r'^([\d\.]+)\s*', img.name)
        if m:
            img_id = m.group(1).rstrip('.').strip()
            image_map[img_id] = img

    common_ids = set(table_map) & set(image_map)
    print(f"Found {len(common_ids)} matched image/table pairs "
          f"({len(image_map)} images, {len(table_map)} tables)")

    skipped = 0
    for sid in tqdm(sorted(common_ids), desc="Generating masks"):
        img_path   = image_map[sid]
        docx_path  = table_map[sid]

        img  = cv2.imread(str(img_path))
        if img is None:
            skipped += 1
            continue
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        h, w = gray.shape

        zone_data = parse_docx(docx_path)
        if not zone_data:
            skipped += 1
            continue

        discs = find_discs(gray)
        if not discs:
            # fallback: use image centre
            discs = [(w//2, h//2, 20)]

        # median disc radius as pixel reference
        median_disc_r = float(np.median([d[2] for d in discs]))
        px_per_mm     = (median_disc_r * 2) / DISC_DIAMETER_MM

        combined_mask = np.zeros((h, w), dtype=np.uint8)

        # pair each zone measurement to its nearest disc
        # (sort discs left-to-right to match typical plate layout)
        discs_sorted = sorted(discs, key=lambda c: c[0])
        n_discs = len(discs_sorted)

        for i, (abbrev, zone_mm) in enumerate(zone_data):
            disc_idx  = min(i, n_discs - 1)
            cx, cy, r = discs_sorted[disc_idx]
            disc_r    = r if r else median_disc_r
            zone_r_px = (zone_mm / 2) * px_per_mm
            draw_zone(combined_mask, cx, cy, zone_r_px, disc_r)

        # save mask using the original image stem so dataset pairing works
        orig_stem = img_path.stem
        mask_file = masks_path / f"{orig_stem}_mask.png"
        cv2.imwrite(str(mask_file), combined_mask)

        # optional debug overlay
        if visualise:
            vis = img.copy()
            overlay = vis.copy()
            overlay[combined_mask > 0] = (0, 200, 80)
            vis = cv2.addWeighted(vis, 0.65, overlay, 0.35, 0)
            for cx, cy, r in discs:
                cv2.circle(vis, (cx, cy), r, (0, 0, 255), 2)
            cv2.imwrite(str(masks_path / f"{orig_stem}_vis.jpg"), vis)

    print(f"\n✅  Masks saved to: {masks_dir}")
    if skipped:
        print(f"⚠️   Skipped {skipped} samples (no image or empty table).")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Generate U-Net masks from Dryad SIRscan docx tables"
    )
    parser.add_argument("--images_dir", required=True,
                        help="Path to images_original folder")
    parser.add_argument("--tables_dir", required=True,
                        help="Path to Tables folder (containing .docx files)")
    parser.add_argument("--masks_dir",  default="data/dryad/masks",
                        help="Output folder for masks")
    parser.add_argument("--visualise",  action="store_true",
                        help="Save colour overlay debug images")
    args = parser.parse_args()
    main(args.images_dir, args.tables_dir, args.masks_dir, args.visualise)
